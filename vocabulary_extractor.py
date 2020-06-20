import docx
import re
import random


DOCUMENT_NAME = "Vocabulary_example.docx"


def choose_random_words(document_name, quantity):
    results = []
    document = docx.Document(document_name)
    while len(results) < quantity:
        index = get_random_index(len(document.paragraphs)-1)
        object = document.paragraphs[index].text
        if matches_criteria(object):
            results.append(object)
    return results


def get_random_index(max_int):
    random_index = random.randint(0, max_int)
    return random_index


def matches_criteria(word):
    """Applies regex to a given word. Return True if word is accepted"""
    prog = re.compile(r"\w+\s?\-\s?\w+\s?\w+?")
    result = prog.search(word)
    return True if result else False

if __name__ == '__main__':
    print(choose_random_words(DOCUMENT_NAME, 2))